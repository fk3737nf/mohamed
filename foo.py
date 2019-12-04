import requests
from openpyxl import Workbook
from docx import Document

url = 'http://api.worldbank.org/v2/country?format=json&per_page=400'

world_bank_data = requests.get(url).json()
list_of_things = world_bank_data[1]

countries = {}

for entry in list_of_things: # looping around countries nad capitals
    country = entry['name']
    capital = entry['capitalCity']

    #this if statment will help print capitals if not it will get rid of it
    if capital:
        countries[country] = capital
# print(countries_capital)
workbook = Workbook()
worksheet = workbook.active
worksheet.title = ('Countries and thier Capitals') #title

# creating worksheet.cell for countries name and countries capital
worksheet.cell(1, 1, 'countriesName')#countriesName
worksheet.cell(1, 2, 'countriesCapital')  # countriescapital

document = Document()

document.add_heading ('Countries and thier Capitals',0) # titl
# looping
row = 2
for country, capital in countries.items():

    worksheet.cell(row, 1, country)
    worksheet.cell(row, 2, capital)

    document.add_paragraph(f'capital of {country} is {capital}') # this wil brig the capital and country on word
    row += 1


document.save('mohamed.docx')
workbook.save('adan.xlsx')